from io import BytesIO
from pathlib import Path
from typing import Dict, Any
import re

from src.core.logging.logger import get_logger

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
except ImportError:
    # Fallback for environments where python-docx is not available
    Document = None
    WD_BREAK = None


logger = get_logger(__name__)


def print(*args, **kwargs):
    message = " ".join(str(arg) for arg in args)
    if message.startswith("Error"):
        logger.error(message)
    elif message.startswith("Warning"):
        logger.warning(message)
    else:
        logger.info(message)


class DocxRenderer:
    def __init__(self, template_name: str = "standard_nttdata", template_style: str = "standard"):
        """
        Initialize DOCX renderer with template style selection.
        
        Args:
            template_name: Template directory name
            template_style: Template style to use
                - "standard": Traditional table-based NTT DATA format (for internal use)
                - "modern": Clean 2026 format with minimal tables (for external clients)
                - "hybrid": Best of both - structured tables for skills, clean format for experience
        """
        self.template_name = template_name
        self.template_style = template_style.lower()
        
        # Validate template style
        if self.template_style not in ["standard", "modern", "hybrid"]:
            print(f"Warning: Invalid template style '{template_style}', defaulting to 'standard'")
            self.template_style = "standard"
        
        # Determine which template file to use based on style
        template_dir = Path("src/templates") / template_name
        
        if self.template_style == "modern":
            # Try to find modern template
            modern_path = template_dir / "template_modern.docx"
            if modern_path.exists():
                self.template_path = modern_path
            else:
                # Fallback to standard template
                print(f"Warning: Modern template not found at {modern_path}, using standard template")
                self.template_style = "standard"
        
        if self.template_style == "hybrid":
            # Try to find hybrid template
            hybrid_path = template_dir / "template_hybrid.docx"
            if hybrid_path.exists():
                self.template_path = hybrid_path
            else:
                # Fallback to standard template
                print(f"Warning: Hybrid template not found at {hybrid_path}, using standard template")
                self.template_style = "standard"
        
        if self.template_style == "standard":
            # Check for both .docx and .doc extensions
            docx_path = template_dir / "template.docx"
            doc_path = template_dir / "template.doc"
            
            if docx_path.exists():
                self.template_path = docx_path
            elif doc_path.exists():
                self.template_path = doc_path
            else:
                # Default to .docx for backward compatibility
                self.template_path = docx_path
        
        # Enhanced field mapping for better placeholder replacement
        # CRITICAL: Field mapping keys MUST match the exact placeholder names in the template
        # Template uses {{phone}}, {{portal_id}}, etc.
        self.field_mapping = {
            # Header placeholders - match template exactly
            'full_name': ['full_name', 'name'],
            'job_title': ['current_title', 'job_title', 'title', 'designation'],
            'email': ['email', 'email_address'],
            'phone': ['contact_number', 'phone', 'mobile', 'contact'],  # Template uses {{phone}}
            'portal_id': ['employee_id', 'portal_id', 'emp_id'],  # Template uses {{portal_id}}
            'location': ['location', 'city'],
            'experience': ['total_experience', 'experience', 'years_of_experience'],
            'professional_summary': ['summary', 'professional_summary'],  # Template uses {{professional_summary}}
            'current_organization': ['organization', 'current_organization', 'company'],
            'target_role': ['target_role', 'desired_position'],
            'grade': ['grade', 'level'],
            
            # Skills table placeholders
            'core_competencies': ['skills', 'primary_skills', 'core_competencies'],  # Template uses {{core_competencies}}
            'technical_skills_section': ['technical_skills_section', 'skills', 'primary_skills'],
            'secondary_skills': ['secondary_skills'],
            'ai_frameworks': ['ai_frameworks'],
            'cloud_platforms': ['cloud_platforms'],
            'databases': ['databases'],
            
            # Education table placeholders
            'degree': ['degree'],
            'institution': ['institution'],
            'year': ['year'],
            'grade': ['grade'],
            
            # Certification table placeholders
            'certification': ['certification'],
            'issuer': ['issuer'],
            'cert_year': ['cert_year'],
            
            # Section placeholders
            'experience_section': ['work_experience_section', 'experience_section', 'work_experience'],
            'projects_section': ['project_section', 'projects_section', 'project_experience'],
            'education_section': ['education_section', 'education'],
            'certifications_section': ['certifications_section', 'certifications'],
            'leadership_section': ['leadership_section'],
        }

    def render(self, context: Dict[str, Any]) -> bytes:
        """Render CV using NTT DATA template with placeholder replacement"""
        if Document is None:
            # Fallback to programmatic generation if python-docx not available
            return self._render_fallback(context)
        
        try:
            # Load the NTT DATA template
            if self.template_path.exists():
                doc = Document(str(self.template_path))
                # Replace placeholders in the template
                self._replace_placeholders(doc, context)
            else:
                # Create from scratch if template not found
                doc = self._create_template_structure(context)

            # Removed page break after header to prevent unnecessary gaps
            # self._add_page_break_after_header(doc, context)

            # Enforce export header/footer formatting across templates.
            self._apply_export_header_footer(doc)
            
            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()
            
        except Exception as e:
            # Fallback to programmatic generation on any error
            print(f"Warning: Template rendering failed ({e}), using fallback")
            return self._render_fallback(context)

    def _add_page_break_after_header(self, doc: Document, context: Dict[str, Any]) -> None:
        """
        Add a page break after the header/personal information section if document is multi-section.
        This ensures clean separation between header and content.
        """
        try:
            if len(doc.paragraphs) < 2:
                return
            
            # Look for header section indicators (name, email, personal info patterns)
            header_end_idx = -1
            found_header = False
            
            for idx, para in enumerate(doc.paragraphs):
                para_text = para.text.strip().lower()
                
                # Check for personal info section indicators
                if any(indicator in para_text for indicator in [
                    "personal", "information", "contact", "details", "profile"
                ]):
                    found_header = True
                
                # Check for main content start indicators
                if found_header and any(section in para_text for section in [
                    "summary", "professional summary", "executive summary",
                    "skills", "expertise", "work experience", "experience",
                    "projects", "project", "education", "qualifications"
                ]):
                    header_end_idx = idx
                    break
            
            # If we found a header section, add a page break
            if header_end_idx > 0 and header_end_idx < len(doc.paragraphs):
                # Insert a page break via a run on the paragraph before the content section
                para = doc.paragraphs[header_end_idx]
                run = para.insert_paragraph_before().add_run()
                run.add_break(WD_BREAK.PAGE)
                
        except Exception as e:
            # Page break is cosmetic, so don't fail the whole export
            print(f"Warning: Failed to add page break after header: {e}")

    def _apply_export_header_footer(self, doc: Document) -> None:
        """Apply standardized header/footer formatting to all sections."""
        if Document is None:
            return

        try:
            logo_path = self._resolve_header_logo_path()

            for section in doc.sections:
                # Header logo: always left aligned.
                self._format_header_with_logo(section.header, logo_path)

                first_page_header = getattr(section, "first_page_header", None)
                if first_page_header is not None:
                    self._format_header_with_logo(first_page_header, logo_path)

                even_page_header = getattr(section, "even_page_header", None)
                if even_page_header is not None:
                    self._format_header_with_logo(even_page_header, logo_path)

                # Footer page numbering: right aligned dynamic fields.
                self._format_footer_with_page_numbers(section.footer)

                first_page_footer = getattr(section, "first_page_footer", None)
                if first_page_footer is not None:
                    self._format_footer_with_page_numbers(first_page_footer)

                even_page_footer = getattr(section, "even_page_footer", None)
                if even_page_footer is not None:
                    self._format_footer_with_page_numbers(even_page_footer)
        except Exception as e:
            print(f"Warning: Could not apply export header/footer formatting: {e}")

    @staticmethod
    def _resolve_header_logo_path() -> Path | None:
        """Resolve the export header logo path relative to repository root."""
        repo_root = Path(__file__).resolve().parents[3]
        logo_path = repo_root / "web-ui" / "static" / "img" / "nttdata_header.png"
        return logo_path if logo_path.exists() else None

    def _format_header_with_logo(self, header, logo_path: Path | None) -> None:
        """Place logo in header, left-aligned."""
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        # Clear existing runs so old text/logo doesn't overlap.
        for run in list(paragraph.runs):
            run.text = ""

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if logo_path is None:
            return

        try:
            run = paragraph.add_run()
            run.add_picture(str(logo_path), width=Inches(1.9))
        except Exception as e:
            print(f"Warning: Failed to place header logo '{logo_path}': {e}")

    def _format_footer_with_page_numbers(self, footer) -> None:
        """Ensure footer has right-aligned dynamic page numbers."""
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        # Remove existing static footer text/runs and rebuild.
        p_element = paragraph._p
        for child in list(p_element):
            p_element.remove(child)

        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        paragraph.add_run("Page ")
        self._append_field(paragraph, "PAGE")
        paragraph.add_run(" of ")
        self._append_field(paragraph, "NUMPAGES")

    @staticmethod
    def _append_field(paragraph, instruction: str) -> None:
        """Append a dynamic Word field (e.g., PAGE, NUMPAGES) to a paragraph."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run = paragraph.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = f" {instruction} "
        run._r.append(instr_text)

        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_separate)

        fallback_text = OxmlElement("w:t")
        fallback_text.text = "1"
        run._r.append(fallback_text)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

    def render_to_file(self, context: Dict[str, Any], output_path: str) -> None:
        """Render CV to a file"""
        try:
            # Generate the document content
            doc_bytes = self.render(context)
            
            # Write to file
            with open(output_path, 'wb') as f:
                f.write(doc_bytes)
                
            print(f"Successfully rendered CV to: {output_path}")
            
        except Exception as e:
            print(f"Error rendering to file: {e}")
            raise

    def _replace_placeholders(self, doc: Document, context: Dict[str, Any]) -> None:
        """Replace {{placeholder}} variables in the document with actual data"""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_paragraph_placeholders(paragraph, context)
        
        # Replace in tables and populate structured data
        for table_idx, table in enumerate(doc.tables):
            # Check if this is a structured data table and populate it
            table_populated = self._populate_structured_table(table, context)
            
            # If not a structured table, do regular paragraph replacement
            if not table_populated:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_paragraph_placeholders(paragraph, context)
        
        # Replace in headers and footers
        for section in doc.sections:
            # Header
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_paragraph_placeholders(paragraph, context)
            
            # Footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                self._replace_paragraph_placeholders(paragraph, context)
        
        # Apply aggressive empty section hiding for ALL templates
        self._hide_empty_sections(doc, context)
        
        # Apply style-specific enhancements
        if self.template_style == "modern":
            self._apply_modern_enhancements(doc, context)
        elif self.template_style == "hybrid":
            self._apply_hybrid_enhancements(doc, context)

    def _replace_paragraph_placeholders(self, paragraph, context: Dict[str, Any]) -> None:
        """Replace placeholders in a single paragraph with enhanced field recognition"""
        # Get the full text of the paragraph
        full_text = paragraph.text
        
        # Find all placeholders in the format {{key}} and also check for static values that match context
        placeholders = re.findall(r'\{\{(\w+)\}\}', full_text)
        
        # Enhanced placeholder replacement
        new_text = full_text
        replaced = False
        
        # Render summary placeholder as ATS-friendly bullet lines.
        if any(ph in {"professional_summary", "summary"} for ph in placeholders):
            summary_value = self._get_mapped_value("professional_summary", context) or self._get_mapped_value("summary", context)
            summary_lines = self._split_summary_lines(summary_value)
            if summary_lines:
                paragraph.clear()
                if len(summary_lines) == 1:
                    paragraph.add_run(summary_lines[0])
                else:
                    for idx, line in enumerate(summary_lines):
                        run = paragraph.add_run(f"• {line}")
                        if idx < len(summary_lines) - 1:
                            run.add_break()
                return

        # Render projects with bold labels for ATS readability.
        project_placeholders = {"projects_section", "project_section", "project_experience"}
        project_ph = next((ph for ph in placeholders if ph in project_placeholders), None)
        if project_ph:
            project_value = self._get_mapped_value(project_ph, context)
            if project_value and str(project_value).strip():
                paragraph.clear()
                self._render_project_lines_with_bold_labels(paragraph, str(project_value))
                return

        # Replace {{placeholder}} format using field mapping
        for placeholder in placeholders:
            placeholder_pattern = f"{{{{{placeholder}}}}}"
            replacement_value = self._get_mapped_value(placeholder, context)
            # Always replace, even if value is empty string (to remove unreplaced placeholders)
            if replacement_value is not None:
                new_text = new_text.replace(placeholder_pattern, replacement_value)
                replaced = True
                print(f"DEBUG: Replaced {{{{ {placeholder} }}}} with '{replacement_value}'")
        
        # Handle the specific template field patterns
        # Pattern 1: NAME:, (with comma)
        if context.get("full_name") and "NAME:," in full_text:
            new_text = full_text.replace("NAME:,", f"NAME: {context['full_name']},")
            replaced = True
        
        # Pattern 2: POTAL ID: (note the typo in template)
        if context.get("employee_id") and "POTAL ID:" in full_text:
            new_text = full_text.replace("POTAL ID:", f"POTAL ID: {context['employee_id']}")
            replaced = True
        
        # Pattern 3: CURRENT GRADE:
        if context.get("grade") and "CURRENT GRADE:" in full_text:
            new_text = full_text.replace("CURRENT GRADE:", f"CURRENT GRADE: {context['grade']}")
            replaced = True
        
        # Pattern 4: Experience Summary - Add content after the header
        if full_text.strip() == "Experience Summary" and context.get("summary"):
            paragraph.clear()
            header_run = paragraph.add_run("Experience Summary")
            header_run.bold = True

            summary_lines = self._split_summary_lines(context.get("summary"))
            if len(summary_lines) <= 1:
                paragraph.add_run("\n\n" + str(context["summary"]))
            else:
                for line in summary_lines:
                    paragraph.add_run("\n• " + line)
            replaced = True
            return  # Early return since we've manually handled this paragraph
        
        # Apply broad regex replacements only for short label-like paragraphs.
        is_label_like = len(full_text.strip()) <= 180 and "\n" not in full_text

        # Also handle generic patterns for broader compatibility
        if is_label_like and context.get("employee_id"):
            # Replace various employee ID patterns - FIXED: Changed (\d*) to ([^\n,]*) to match actual content
            emp_id_patterns = [
                r'(Portal\s*id[/\s]*emp\s*id\s*:?\s*)([^\n,]*)',
                r'(Employee\s*ID\s*:?\s*)([^\n,]*)',
                r'(EMP\s*ID\s*:?\s*)([^\n,]*)',
                r'(Portal\s*ID\s*:?\s*)([^\n,]*)',
                r'(POTAL\s*ID\s*:?\s*)([^\n,]*)',  # Handle template typo
            ]
            
            for pattern in emp_id_patterns:
                if re.search(pattern, new_text, re.IGNORECASE):
                    new_text = re.sub(pattern, f'\\g<1>{context["employee_id"]}', new_text, flags=re.IGNORECASE)
                    replaced = True
        
        if is_label_like and context.get("contact_number"):
            # Replace contact number patterns - FIXED: Changed (\d*) to ([^\n,]*) to match actual content
            contact_patterns = [
                r'(Contact\s*Details?\s*:?\s*)([^\n,]*)',
                r'(Phone\s*:?\s*)([^\n,]*)',
                r'(Mobile\s*:?\s*)([^\n,]*)',
                r'(Contact\s*:?\s*)([^\n,]*)',
            ]
            
            for pattern in contact_patterns:
                if re.search(pattern, new_text, re.IGNORECASE):
                    new_text = re.sub(pattern, f'\\g<1>{context["contact_number"]}', new_text, flags=re.IGNORECASE)
                    replaced = True
        
        # Replace other common field patterns - ENHANCED with better regex patterns
        field_patterns = {
            # Match only dedicated Name lines; avoid replacing substrings like
            # "Project Name:" inside content sections.
            "full_name": [r'^\s*(Name\s*:?,?\s*)([^\n,]*)$'],
            "email": [r'(Email\s*:?\s*)([a-zA-Z0-9._%+-]*@?[a-zA-Z0-9.-]*\.?[a-zA-Z]*)'],
            "current_title": [r'(Title\s*:?\s*)([A-Za-z\s]*)', r'(Designation\s*:?\s*)([A-Za-z\s]*)'],
            "location": [r'(Location\s*:?\s*)([A-Za-z\s,.-]*)'],  # Enhanced to match commas, periods, hyphens
            "organization": [r'(Current\s*Organization\s*:?\s*)([A-Za-z\s]*)', r'(Organization\s*:?\s*)([A-Za-z\s]*)', r'(Company\s*:?\s*)([A-Za-z\s]*)'],  # Added "Current Organization"
            # Require an explicit colon to avoid matching headings like
            # "PROFESSIONAL EXPERIENCE" and appending numeric values there.
            "experience": [r'(Experience\s*:\s*)([^\n,]*)'],
            "grade": [r'(Grade\s*:?\s*)([A-Za-z0-9\s]*)', r'(CURRENT\s*GRADE\s*:?\s*)([A-Za-z0-9\s]*)'],
        }
        
        if is_label_like:
            for field_key, patterns in field_patterns.items():
                field_value = context.get(field_key)
                if field_value:
                    for pattern in patterns:
                        if re.search(pattern, new_text, re.IGNORECASE):
                            new_text = re.sub(pattern, f'\\g<1>{field_value}', new_text, flags=re.IGNORECASE)
                            replaced = True
        
        # Update paragraph text while preserving formatting
        if replaced and new_text != full_text:
            # Clear existing runs
            paragraph.clear()
            # Add new text
            paragraph.add_run(new_text)

    def _render_project_lines_with_bold_labels(self, paragraph, project_text: str) -> None:
        """Render project section lines with selected labels in bold."""
        lines = str(project_text or "").replace("\r\n", "\n").split("\n")
        bold_labels = ["Project Name:", "Client:", "Role:", "Duration:", "Description:"]
        label_pattern = re.compile(r"^(Project Name:|Client:|Role:|Duration:|Description:)\s*(.*)$", re.IGNORECASE)

        wrote_any = False
        for idx, raw_line in enumerate(lines):
            line = raw_line.strip()
            if not line:
                if wrote_any and idx < len(lines) - 1:
                    paragraph.add_run("").add_break()
                continue

            m = label_pattern.match(line)
            if m:
                label = next((lbl for lbl in bold_labels if lbl.lower() == m.group(1).lower()), m.group(1))
                value = m.group(2).strip()
                label_run = paragraph.add_run(label)
                label_run.bold = True
                if value:
                    paragraph.add_run(f" {value}")
            else:
                paragraph.add_run(line)

            wrote_any = True
            if idx < len(lines) - 1:
                paragraph.add_run("").add_break()
    
    def _get_mapped_value(self, placeholder: str, context: Dict[str, Any]) -> str:
        """Get value for placeholder using field mapping with smart formatting"""
        # First try direct match in context
        if placeholder in context:
            value = context.get(placeholder)
            if value is not None:
                # Format complex data structures appropriately
                formatted = self._format_value_for_placeholder(placeholder, value)
                if formatted and formatted.strip():
                    return formatted.strip()
                # Return empty string even if value is empty (to replace placeholder)
                return ""
        
        # Try field mapping - check if any mapped field has a value
        if placeholder in self.field_mapping:
            for field_key in self.field_mapping[placeholder]:
                value = context.get(field_key)
                if value is not None:
                    formatted = self._format_value_for_placeholder(placeholder, value)
                    if formatted and formatted.strip():
                        return formatted.strip()
        
        # If not found anywhere, return empty string to remove placeholder
        # Debug: Log if placeholder not found
        print(f"DEBUG: Placeholder '{placeholder}' not found, replacing with empty string")
        
        return ""
    
    def _format_value_for_placeholder(self, placeholder: str, value: Any) -> str:
        """Format value appropriately based on placeholder type and value structure"""
        # Handle None
        if value is None:
            return ""
        
        # Handle simple strings
        if isinstance(value, str):
            return value
        
        # Handle simple types (int, float, bool)
        if isinstance(value, (int, float, bool)):
            return str(value)
        
        # Handle complex data structures based on placeholder type
        placeholder_lower = placeholder.lower()
        
        # Project-related placeholders
        if 'project' in placeholder_lower and isinstance(value, list):
            return self._format_projects_for_text(value)
        
        # Education-related placeholders
        if 'education' in placeholder_lower and isinstance(value, list):
            return self._format_education_for_text(value)
        
        # Certification-related placeholders
        if 'certification' in placeholder_lower and isinstance(value, (list, str)):
            if isinstance(value, list):
                return self._format_certifications_for_text(value)
            return value
        
        # Experience-related placeholders
        if 'experience' in placeholder_lower and isinstance(value, list):
            return self._format_experience_for_text(value)
        
        # Skills-related - handle lists
        if 'skill' in placeholder_lower and isinstance(value, list):
            return ', '.join(str(s) for s in value if s)
        
        # For any other list, convert to comma-separated
        if isinstance(value, list):
            # Filter out None and empty strings
            filtered = [str(item) for item in value if item]
            return ', '.join(filtered) if filtered else ""
        
        # For dictionaries, return empty (shouldn't be used for simple placeholders)
        if isinstance(value, dict):
            return ""
        
        # Default: convert to string
        return str(value)
    
    def _format_projects_for_text(self, projects: list) -> str:
        """Format projects list as clean text without markdown"""
        if not projects:
            return ""
        
        formatted_sections = []
        
        for proj in projects:
            if isinstance(proj, dict):
                lines = []
                
                # Project Name (uppercase, no markdown)
                project_name = proj.get("project_name", proj.get("name", ""))
                if project_name:
                    lines.append(project_name.upper())
                    lines.append("")  # Blank line
                
                # Client
                client = proj.get("client", "")
                if client:
                    lines.append(f"Client: {client}")
                
                # Description (remove markdown syntax)
                description = proj.get("project_description", proj.get("description", ""))
                if description:
                    # Remove markdown bold syntax
                    description = re.sub(r'\*\*([^*]+)\*\*', r'\1', description)
                    # Extract just the description part if it contains "Client:" and "Description:" markers
                    if "Description:" in description:
                        desc_match = re.search(r'Description:\s*(.+)', description, re.IGNORECASE)
                        if desc_match:
                            description = desc_match.group(1).strip()
                    lines.append(f"Description: {description}")
                
                # Duration
                duration = proj.get("duration", "")
                if duration:
                    lines.append(f"Duration: {duration}")
                
                # Role
                role = proj.get("role", proj.get("position", ""))
                if role:
                    lines.append(f"Role: {role}")
                
                # Technologies
                technologies = self._format_technologies(
                    proj.get("technologies_used", 
                    proj.get("technologies", 
                    proj.get("tech_stack", [])))
                )
                if technologies:
                    lines.append(f"Technologies: {technologies}")
                
                # Responsibilities
                responsibilities = proj.get("responsibilities", 
                                           proj.get("key_responsibilities", 
                                           proj.get("contributions", [])))
                if responsibilities:
                    lines.append("Roles and Responsibilities:")
                    if isinstance(responsibilities, list):
                        for resp in responsibilities:
                            lines.append(f"  • {resp}")
                    elif isinstance(responsibilities, str):
                        lines.append(f"  • {responsibilities}")
                
                formatted_sections.append("\n".join(lines))
        
        return "\n\n".join(formatted_sections)
    
    def _format_education_for_text(self, education: list) -> str:
        """Format education list as clean text"""
        if not education:
            return ""
        
        formatted_items = []
        
        for edu in education:
            if isinstance(edu, dict):
                parts = []
                
                # Degree and specialization
                qualification = edu.get("qualification", "")
                specialization = edu.get("specialization", "")
                if qualification:
                    if specialization:
                        parts.append(f"{qualification} in {specialization}")
                    else:
                        parts.append(qualification)
                
                # College/University
                college = edu.get("college", "")
                university = edu.get("university", "")
                if college:
                    parts.append(f"from {college}")
                elif university:
                    parts.append(f"from {university}")
                
                # Year
                year = edu.get("year", "")
                if year:
                    parts.append(f"({year})")
                
                # Percentage/CGPA
                percentage = edu.get("percentage", "")
                if percentage:
                    parts.append(f"- {percentage}")
                
                if parts:
                    formatted_items.append(" ".join(parts))
        
        return "\n• ".join([""] + formatted_items) if formatted_items else ""
    
    def _format_certifications_for_text(self, certifications: list) -> str:
        """Format certifications list as clean text"""
        if not certifications:
            return ""
        
        formatted_items = []
        
        for cert in certifications:
            if isinstance(cert, dict):
                name = cert.get("name", cert.get("certification", ""))
                year = cert.get("year", cert.get("date", ""))
                issuer = cert.get("issuer", "")
                
                cert_text = name
                if issuer:
                    cert_text += f" - {issuer}"
                if year:
                    cert_text += f" ({year})"
                
                if cert_text:
                    formatted_items.append(cert_text)
            elif isinstance(cert, str):
                formatted_items.append(cert)
        
        return "\n• ".join([""] + formatted_items) if formatted_items else ""
    
    def _format_experience_for_text(self, experience: list) -> str:
        """Format work experience list as clean text"""
        if not experience:
            return ""
        
        formatted_sections = []
        
        for exp in experience:
            if isinstance(exp, dict):
                lines = []
                
                # Company & Title
                company = exp.get("company", exp.get("organization", ""))
                title = exp.get("title", exp.get("designation", exp.get("role", "")))
                
                if title and company:
                    lines.append(f"{title} at {company}")
                elif title:
                    lines.append(title)
                elif company:
                    lines.append(company)
                
                # Duration
                duration = exp.get("duration", "")
                if duration:
                    lines.append(f"Duration: {duration}")
                
                # Location
                location = exp.get("location", "")
                if location:
                    lines.append(f"Location: {location}")
                
                # Responsibilities
                responsibilities = exp.get("responsibilities", exp.get("achievements", []))
                if responsibilities:
                    lines.append("Key Responsibilities:")
                    if isinstance(responsibilities, list):
                        for resp in responsibilities:
                            lines.append(f"  • {resp}")
                    elif isinstance(responsibilities, str):
                        lines.append(f"  • {responsibilities}")
                
                if lines:
                    formatted_sections.append("\n".join(lines))
        
        return "\n\n".join(formatted_sections)

    def _create_template_structure(self, context: Dict[str, Any]) -> Document:
        """Create NTT DATA-style document structure when template is not available"""
        doc = Document()
        
        # Add NTT DATA header styling
        header_section = doc.sections[0]
        header = header_section.header
        header_para = header.paragraphs[0]
        header_para.text = "NTT DATA"
        
        # Document title
        title = doc.add_heading("Professional CV", level=1)
        
        # Personal Information Section
        doc.add_heading("Personal Information", level=2)
        
        # Create structured layout similar to NTT template
        personal_info = [
            ("Name", context.get("full_name", "")),
            ("Employee ID", context.get("employee_id", "")),
            ("Email", context.get("email", "")),
            ("Contact Number", context.get("contact_number", "")),
            ("Current Title", context.get("current_title", "")),
            ("Grade", context.get("grade", "")),
            ("Location", context.get("location", "")),
            ("Organization", context.get("organization", "")),
            ("Total Experience", context.get("experience", "")),
            ("Target Role", context.get("target_role", "")),
        ]
        
        # Add personal info as table for better formatting
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Table Grid'
        
        for label, value in personal_info:
            if value:  # Only add rows with data
                row_cells = info_table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = str(value)
        
        # Professional Summary
        if context.get("summary"):
            doc.add_heading("Professional Summary", level=2)
            summary_lines = self._split_summary_lines(context.get("summary"))
            if len(summary_lines) <= 1:
                doc.add_paragraph(context.get("summary"))
            else:
                for line in summary_lines:
                    try:
                        doc.add_paragraph(line, style='List Bullet')
                    except Exception:
                        doc.add_paragraph(f"• {line}")
        
        # Skills sections
        self._add_skills_sections(doc, context)
        
        # Experience sections
        self._add_experience_sections(doc, context)
        
        # Education and Additional sections
        self._add_education_sections(doc, context)
        
        return doc

    def _add_skills_sections(self, doc: Document, context: Dict[str, Any]) -> None:
        """Add skills sections to the document"""
        skills_sections = [
            ("Primary Skills", "skills"),
            ("Secondary Skills", "secondary_skills"),
            ("Tools & Platforms", "tools_and_platforms"),
            ("AI Frameworks", "ai_frameworks"),
            ("Cloud Platforms", "cloud_platforms"),
            ("Operating Systems", "operating_systems"),
            ("Databases", "databases"),
            ("Domain Expertise", "domain_expertise"),
        ]
        
        for section_title, context_key in skills_sections:
            content = context.get(context_key, "")
            if content:
                doc.add_heading(section_title, level=2)
                doc.add_paragraph(str(content))

    def _add_experience_sections(self, doc: Document, context: Dict[str, Any]) -> None:
        """Add experience sections to the document"""
        # Work Experience
        if context.get("experience_section"):
            doc.add_heading("Work Experience", level=2)
            exp_content = str(context.get("experience_section", ""))
            for line in exp_content.split("\n\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
        elif context.get("work_experience"):
            doc.add_heading("Work Experience", level=2)
            doc.add_paragraph(str(context.get("work_experience")))
        
        # Project Experience
        if context.get("projects_section"):
            doc.add_heading("Project Experience", level=2)
            proj_content = str(context.get("projects_section", ""))
            for line in proj_content.split("\n\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
        elif context.get("project_experience"):
            doc.add_heading("Project Experience", level=2)
            doc.add_paragraph(str(context.get("project_experience")))
        
        # Leadership & Impact
        if context.get("leadership_section"):
            doc.add_heading("Leadership Experience", level=2)
            lead_content = str(context.get("leadership_section", ""))
            for line in lead_content.split("\n\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
        elif context.get("leadership_lines") and isinstance(context.get("leadership_lines"), list) and context.get("leadership_lines"):
            doc.add_heading("Leadership Experience", level=2)
            for lead in context.get("leadership_lines", []):
                if isinstance(lead, dict):
                    title = lead.get("title", "")
                    organization = lead.get("organization", "")
                    if title or organization:
                        doc.add_paragraph(f"{title} at {organization}", style='List Number')
                        team_size = lead.get("team_size", "")
                        if team_size:
                            doc.add_paragraph(f"Team Size: {team_size}")
                        responsibilities = lead.get("responsibilities", [])
                        if responsibilities:
                            for resp in responsibilities:
                                doc.add_paragraph(str(resp), style='List Bullet 2')

    def _add_education_sections(self, doc: Document, context: Dict[str, Any]) -> None:
        """Add education and additional sections"""
        # Education - use formatted section if available
        if context.get("education_section"):
            doc.add_heading("Education", level=2)
            edu_content = str(context.get("education_section", ""))
            for line in edu_content.split("\n\n"):
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    # Apply bullet formatting if it contains multiple lines
                    if "\n" in line:
                        para.style = 'List Bullet'
        elif context.get("education"):
            doc.add_heading("Education", level=2)
            education = context.get("education", [])
            if isinstance(education, list):
                for edu in education:
                    if isinstance(edu, dict):
                        degree = edu.get("degree", "") or edu.get("qualification", "")
                        field = edu.get("field_of_study", "") or edu.get("specialization", "")
                        institution = edu.get("institution", "") or edu.get("college", "") or edu.get("university", "")
                        year = edu.get("year", "") or edu.get("graduation_year", "") or edu.get("yearOfPassing", "")
                        grade = edu.get("grade", "") or edu.get("percentage", "") or edu.get("cgpa", "")
                        
                        edu_text = []
                        if degree:
                            if field:
                                edu_text.append(f"{degree} in {field}")
                            else:
                                edu_text.append(degree)
                        if institution:
                            edu_text.append(institution)
                        if year:
                            edu_text.append(f"({year})")
                        if grade:
                            edu_text.append(f"Grade: {grade}")
                        
                        if edu_text:
                            doc.add_paragraph(" | ".join(edu_text))
            else:
                doc.add_paragraph(str(education))
        
        # Certifications
        if context.get("certifications_section"):
            doc.add_heading("Certifications", level=2)
            cert_content = str(context.get("certifications_section", ""))
            for line in cert_content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip(), style='List Bullet' if "•" in line else None)
        elif context.get("certifications"):
            doc.add_heading("Certifications", level=2)
            doc.add_paragraph(str(context.get("certifications")))
        
        # Languages
        languages_section = context.get("languages")
        if languages_section:
            doc.add_heading("Languages", level=2)
            if isinstance(languages_section, list):
                for lang in languages_section:
                    doc.add_paragraph(str(lang), style='List Bullet')
            else:
                doc.add_paragraph(str(languages_section))
        
        # Awards & Recognition
        if context.get("awards"):
            doc.add_heading("Awards & Recognition", level=2)
            awards = context.get("awards", [])
            if isinstance(awards, list):
                for award in awards:
                    doc.add_paragraph(str(award), style='List Bullet')
            else:
                doc.add_paragraph(str(awards))
        
        # Publications
        if context.get("publications"):
            doc.add_heading("Publications", level=2)
            publications = context.get("publications", [])
            if isinstance(publications, list):
                for pub in publications:
                    doc.add_paragraph(str(pub), style='List Bullet')
            else:
                doc.add_paragraph(str(publications))

    def _populate_structured_table(self, table, context: Dict[str, Any]) -> bool:
        """Populate structured data tables based on their content and headers"""
        try:
            # Get table headers from first row to identify table type
            if len(table.rows) == 0:
                return False
                
            header_row = table.rows[0]
            header_text = " ".join([cell.text.strip() for cell in header_row.cells]).lower()
            
            # Technical Expertise Table - ALWAYS KEEP (good for scanning)
            if "elements" in header_text and "particulars" in header_text:
                return self._populate_technical_expertise_table(table, context)
            
            # Project Details Table - Skip for hybrid/modern templates
            elif "project name" in header_text and "client" in header_text:
                if self.template_style in ["hybrid", "modern"]:
                    # For hybrid/modern: Replace table with clean formatted text
                    return self._replace_table_with_clean_projects(table, context)
                else:
                    # For standard: Use traditional table format
                    return self._populate_project_details_table(table, context)
            
            # Experience Details Table - Skip for hybrid/modern templates
            elif "organization" in header_text and "designation" in header_text and "joining" in header_text:
                if self.template_style in ["hybrid", "modern"]:
                    # For hybrid/modern: Replace table with timeline format
                    return self._replace_table_with_clean_experience(table, context)
                else:
                    # For standard: Use traditional table format
                    return self._populate_experience_details_table(table, context)
            
            # Training/Certifications Table - Simplified for hybrid/modern
            elif "course details" in header_text and "duration" in header_text:
                if self.template_style in ["hybrid", "modern"]:
                    return self._replace_table_with_bullet_certifications(table, context)
                else:
                    return self._populate_training_table(table, context)
            
            # Qualification Details Table - KEEP SIMPLE TABLE (expected format)
            elif "degree" in header_text and "college" in header_text and "university" in header_text:
                return self._populate_qualification_table(table, context)
                
            return False
            
        except Exception as e:
            print(f"Error populating structured table: {e}")
            return False

    def _populate_technical_expertise_table(self, table, context: Dict[str, Any]) -> bool:
        """Populate the Technical Expertise table with separate rows for primary and secondary skills"""
        try:
            # Prepare skill categories as separate rows
            skill_rows = []
            
            # Primary Skills (required)
            primary_skills = context.get("skills", "")
            if primary_skills:
                if isinstance(primary_skills, list):
                    primary_str = ", ".join(str(s).strip() for s in primary_skills if s)
                elif isinstance(primary_skills, str):
                    primary_str = primary_skills.strip()
                else:
                    primary_str = ""
                    
                if primary_str:
                    skill_rows.append(("Primary Skills", primary_str))
            
            # Secondary Skills (separate row)
            secondary_skills = context.get("secondary_skills", "")
            if secondary_skills:
                if isinstance(secondary_skills, list):
                    secondary_str = ", ".join(str(s).strip() for s in secondary_skills if s)
                elif isinstance(secondary_skills, str):
                    secondary_str = secondary_skills.strip()
                else:
                    secondary_str = ""
                    
                if secondary_str:
                    skill_rows.append(("Secondary Skills", secondary_str))
            
            # Other technical categories (consolidated into additional rows if needed)
            other_categories = [
                ("AI/ML Frameworks", context.get("ai_frameworks", "")),
                ("Cloud Platforms", context.get("cloud_platforms", "")),
                ("Tools & Platforms", context.get("tools_and_platforms", "")),
                ("Databases", context.get("databases", "")),
                ("Operating Systems", context.get("operating_systems", "")),
                ("Domain Expertise", context.get("domain_expertise", "")),
            ]
            
            for category_name, category_value in other_categories:
                if category_value:
                    if isinstance(category_value, list):
                        category_str = ", ".join(str(s).strip() for s in category_value if s)
                    elif isinstance(category_value, str):
                        category_str = category_value.strip()
                    else:
                        category_str = ""
                    
                    if category_str:
                        skill_rows.append((category_name, category_str))
            
            if not skill_rows:
                return False
            
            # Ensure table has enough rows
            while len(table.rows) < len(skill_rows) + 1:  # +1 for header
                table.add_row()
            
            # Populate each row
            for idx, (label, skills) in enumerate(skill_rows):
                row_idx = idx + 1  # Skip header row
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    if len(row.cells) >= 2:
                        # Set label in first column
                        row.cells[0].paragraphs[0].clear()
                        row.cells[0].paragraphs[0].add_run(label)
                        
                        # Set skills in second column
                        row.cells[1].paragraphs[0].clear()
                        row.cells[1].paragraphs[0].add_run(skills)
            
            # Remove extra unused rows
            while len(table.rows) > len(skill_rows) + 1:
                tbl = table._element
                tbl.remove(table.rows[-1]._element)
            
            print(f"DEBUG: Populated technical expertise table with {len(skill_rows)} skill categories")
            return True
            
        except Exception as e:
            print(f"Error populating technical expertise table: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _populate_project_details_table(self, table, context: Dict[str, Any]) -> bool:
        """Populate the Project Details table with project experience"""
        try:
            # Get project experience data - handle both string and list formats
            project_exp = context.get("project_experience", "")
            projects = []
            
            if isinstance(project_exp, str):
                # Parse project experience if it's formatted text
                projects = self._parse_project_experience(project_exp)
            elif isinstance(project_exp, list):
                # If it's already a list, use it directly but ensure proper field mapping
                for proj in project_exp:
                    if isinstance(proj, dict):
                        # Enhanced field mapping for voice extraction data structure
                        mapped_project = {
                            "name": proj.get("project_name", proj.get("name", "")),
                            "client": proj.get("client", ""),
                            "description": proj.get("project_description", proj.get("description", "")),
                            "technologies": self._format_technologies(
                                proj.get("technologies_used", 
                                proj.get("technologies", 
                                proj.get("tech_stack", [])))
                            ),
                            "duration": proj.get("duration", ""),
                            "role": proj.get("role", proj.get("position", "")),
                            "contributions": self._format_responsibilities(
                                proj.get("responsibilities", 
                                proj.get("key_responsibilities", 
                                proj.get("contributions", [])))
                            ),
                            "team_size": proj.get("team_size", "")
                        }
                        projects.append(mapped_project)
            
            if not projects:
                print("DEBUG: No projects found in project_experience")
                return False
            
            # Debug: Print project mapping results
            print(f"DEBUG: Successfully mapped {len(projects)} projects for DOCX rendering")
            for i, proj in enumerate(projects):
                print(f"  Project {i+1}: '{proj.get('name', 'NO NAME')}' - Client: '{proj.get('client', 'NO CLIENT')}'")
            
            # Populate each project
            row_idx = 1  # Skip header row
            for project in projects:
                if row_idx >= len(table.rows):
                    table.add_row()
                    
                row = table.rows[row_idx]
                cells = row.cells
                
                # Debug: Print cell count and project name
                print(f"DEBUG: Table has {len(cells)} columns, populating project: {project.get('name', 'Unknown')}")
                
                # Map project data to table columns based on template structure
                if len(cells) >= 10:  # Adjusted based on actual template structure
                    # NTT DATA Info (column 0)
                    cells[0].paragraphs[0].clear()
                    cells[0].paragraphs[0].add_run(context.get("organization", "NTT DATA"))
                    
                    # Project Name (column 1) - This was getting "John Smith" before
                    cells[1].paragraphs[0].clear()
                    project_name = project.get("name", "")
                    if project_name:
                        cells[1].paragraphs[0].add_run(project_name)
                        print(f"DEBUG: Set project name to: {project_name}")
                    
                    # Client (column 2)
                    cells[2].paragraphs[0].clear()
                    client = project.get("client", "")
                    if client:
                        cells[2].paragraphs[0].add_run(client)
                    
                    # Project Description (column 3)
                    cells[3].paragraphs[0].clear()
                    description = project.get("description", "")
                    if description:
                        # Truncate long descriptions for table display
                        if len(description) > 100:
                            description = description[:97] + "..."
                        cells[3].paragraphs[0].add_run(description)
                    
                    # Environment (column 4)
                    cells[4].paragraphs[0].clear()
                    technologies = project.get("technologies", "")
                    if technologies:
                        cells[4].paragraphs[0].add_run(technologies)
                    
                    # Duration From (column 5)
                    cells[5].paragraphs[0].clear()
                    duration = project.get("duration", "")
                    if duration:
                        # Split duration and take first part
                        duration_parts = duration.replace(" to ", " - ").split(" - ")
                        if len(duration_parts) >= 1:
                            # Extract month/year from date
                            start_date = duration_parts[0].strip()
                            # Convert "Mar 2023" to "03/23" format
                            formatted_start = self._format_date_for_table(start_date)
                            cells[5].paragraphs[0].add_run(formatted_start)
                    
                    # Duration To (column 6)
                    cells[6].paragraphs[0].clear()
                    if duration and " - " in duration:
                        duration_parts = duration.replace(" to ", " - ").split(" - ")
                        if len(duration_parts) >= 2:
                            end_date = duration_parts[1].strip()
                            formatted_end = self._format_date_for_table(end_date)
                            cells[6].paragraphs[0].add_run(formatted_end)
                    
                    # Role/Responsibility (column 7)
                    cells[7].paragraphs[0].clear()
                    role = project.get("role", "")
                    if role:
                        cells[7].paragraphs[0].add_run(role)
                    
                    # Contributions (column 8) - if available
                    if len(cells) > 8:
                        cells[8].paragraphs[0].clear()
                        contributions = project.get("contributions", "")
                        if contributions:
                            # Truncate for table display
                            if len(contributions) > 80:
                                contributions = contributions[:77] + "..."
                            cells[8].paragraphs[0].add_run(contributions)
                    
                    # Team Size (column 9) - if available
                    if len(cells) > 9:
                        cells[9].paragraphs[0].clear()
                        team_size = project.get("team_size", "")
                        if team_size:
                            cells[9].paragraphs[0].add_run(str(team_size))
                
                row_idx += 1
                
            print(f"DEBUG: Successfully populated {len(projects)} projects")
            return True
            
        except Exception as e:
            print(f"Error populating project details table: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _populate_experience_details_table(self, table, context: Dict[str, Any]) -> bool:
        """Populate the Experience Details table with work history"""
        try:
            # Get work experience data
            work_exp = context.get("work_experience", "")
            if not work_exp:
                return False

            # Parse work experience from either list or text format
            experiences = self._parse_work_experience(work_exp)
            
            # Populate each experience
            row_idx = 1  # Skip header row
            for idx, exp in enumerate(experiences):
                if row_idx >= len(table.rows):
                    table.add_row()
                    
                row = table.rows[row_idx]
                cells = row.cells
                
                if len(cells) >= 5:  # Based on template structure
                    # Sl. No. (column 0)
                    cells[0].paragraphs[0].clear()
                    cells[0].paragraphs[0].add_run(str(idx + 1))
                    
                    # Organization (column 1)
                    cells[1].paragraphs[0].clear()
                    cells[1].paragraphs[0].add_run(exp.get("company", ""))
                    
                    # Designation (column 2)
                    cells[2].paragraphs[0].clear()
                    cells[2].paragraphs[0].add_run(exp.get("title", ""))
                    
                    # Joining Date (column 3)
                    cells[3].paragraphs[0].clear()
                    start_val = exp.get("startDate") or ""
                    end_val = exp.get("endDate") or ""
                    if not start_val:
                        duration_parts = str(exp.get("duration", "")).split(" - ")
                        if len(duration_parts) >= 1:
                            start_val = duration_parts[0].strip()
                        if len(duration_parts) >= 2 and not end_val:
                            end_val = duration_parts[1].strip()
                    if start_val:
                        cells[3].paragraphs[0].add_run(str(start_val))
                    
                    # Relieving Date (column 4)
                    cells[4].paragraphs[0].clear()
                    if end_val:
                        cells[4].paragraphs[0].add_run(str(end_val))
                    
                row_idx += 1
                
            return True
            
        except Exception as e:
            print(f"Error populating experience details table: {e}")
            return False

    def _populate_training_table(self, table, context: Dict[str, Any]) -> bool:
        """Populate the Training/Certifications table"""
        try:
            certifications = context.get("certifications", "")
            if not certifications:
                return False
                
            # Parse certifications
            certs = self._parse_certifications(certifications)
            
            row_idx = 1  # Skip header
            for cert in certs:
                if row_idx >= len(table.rows):
                    table.add_row()
                    
                row = table.rows[row_idx]
                cells = row.cells
                
                if len(cells) >= 4:
                    # Course Details
                    cells[0].paragraphs[0].clear()
                    cells[0].paragraphs[0].add_run(cert.get("name", ""))
                    
                    # Duration (if available)
                    cells[1].paragraphs[0].clear()
                    cells[1].paragraphs[0].add_run(cert.get("duration", ""))
                    
                    # From date
                    cells[2].paragraphs[0].clear()
                    cells[2].paragraphs[0].add_run(cert.get("year", ""))
                    
                    # To date
                    cells[3].paragraphs[0].clear()
                    cells[3].paragraphs[0].add_run(cert.get("year", ""))
                    
                row_idx += 1
                
            return True
            
        except Exception as e:
            print(f"Error populating training table: {e}")
            return False

    def _populate_qualification_table(self, table, context: Dict[str, Any]) -> bool:
        """Populate the Qualification Details table - handles both placeholder replacement and row addition"""
        try:
            education = context.get("education", "")
            if not education:
                # No education data - try placeholder replacement instead
                return False
            
            # Handle both structured list and text format
            educations = []
            if isinstance(education, list):
                # Already structured - use directly
                educations = education
            else:
                # Parse from text format
                educations = self._parse_education(education)
            
            if not educations:
                return False
            
            # Determine table structure from headers
            if len(table.rows) == 0:
                return False
            
            header_row = table.rows[0]
            num_cols = len(header_row.cells)
            
            print(f"DEBUG: Populating education table with {len(educations)} entries, {num_cols} columns")
            
            # Check if row 1 contains placeholders - if so, REPLACE them in the SAME row (don't return False)
            if len(table.rows) >= 2:
                first_data_row = table.rows[1]
                first_row_text = " ".join([cell.text for cell in first_data_row.cells]).lower()
                
                # If row contains placeholders, REPLACE them directly in this method
                if "{{" in first_row_text and "}}" in first_row_text:
                    print("DEBUG: Found placeholders in education table row, replacing them directly")
                    # Replace placeholders in the existing row instead of returning False
                    # This ensures the table replacement happens
                    if len(educations) > 0:
                        edu = educations[0]  # Use first education entry
                        if isinstance(edu, dict):
                            # Extract all fields
                            degree = edu.get("degree") or edu.get("qualification") or edu.get("program") or ""
                            specialization = edu.get("specialization") or edu.get("field_of_study") or edu.get("major") or ""
                            year = edu.get("year") or edu.get("year_of_completion") or edu.get("graduation_year") or ""
                            college = edu.get("college") or edu.get("institution") or ""
                            university = edu.get("university") or edu.get("institution") or ""
                            grade = edu.get("grade") or edu.get("percentage") or edu.get("gpa") or edu.get("cgpa") or ""
                            
                            # Replace placeholders in each cell
                            for cell in first_data_row.cells:
                                for paragraph in cell.paragraphs:
                                    cell_text = paragraph.text
                                    if "{{" in cell_text and "}}" in cell_text:
                                        # Replace all possible placeholders
                                        cell_text = cell_text.replace("{{degree}}", degree)
                                        cell_text = cell_text.replace("{{qualification}}", degree)
                                        cell_text = cell_text.replace("{{specialization}}", specialization)
                                        cell_text = cell_text.replace("{{year}}", str(year))
                                        cell_text = cell_text.replace("{{college}}", college)
                                        cell_text = cell_text.replace("{{institution}}", university or college)  # Use university if available
                                        cell_text = cell_text.replace("{{university}}", university)
                                        cell_text = cell_text.replace("{{grade}}", str(grade))
                                        cell_text = cell_text.replace("{{percentage}}", str(grade))
                                        cell_text = cell_text.replace("{{gpa}}", str(grade))
                                        cell_text = cell_text.replace("{{cgpa}}", str(grade))
                                        
                                        # Update the cell
                                        paragraph.clear()
                                        paragraph.add_run(cell_text)
                                        print(f"DEBUG: Replaced education placeholder in cell: {cell_text}")
                    
                    return True  # Return True to indicate we handled the table
            
            # Otherwise use row addition strategy
            row_idx = 1  # Skip header
            for idx, edu in enumerate(educations):
                if not isinstance(edu, dict):
                    continue
                
                if row_idx >= len(table.rows):
                    table.add_row()
                    
                row = table.rows[row_idx]
                cells = row.cells
                
                # Enhanced field extraction with multiple fallback names
                degree = edu.get("degree") or edu.get("qualification") or edu.get("program") or ""
                specialization = edu.get("specialization") or edu.get("field_of_study") or edu.get("major") or ""
                year = edu.get("year") or edu.get("year_of_completion") or edu.get("graduation_year") or ""
                college = edu.get("college") or edu.get("institution") or ""
                university = edu.get("university") or edu.get("institution") or ""
                grade = edu.get("grade") or edu.get("percentage") or edu.get("gpa") or edu.get("cgpa") or ""
                
                # Clean up university field if it contains newlines or extra text
                if university:
                    # Remove extra whitespace and newlines
                    university = " ".join(university.split())
                    # If it's too long (combined with other text), try to extract meaningful part
                    if len(university) > 100:
                        # Try to find "University" keyword and extract around it
                        if "University" in university:
                            parts = university.split("University")
                            if len(parts) >= 2:
                                # Take the part before "University" and include "University"
                                university = parts[0].strip() + " University"
                        # Truncate if still too long
                        if len(university) > 80:
                            university = university[:77] + "..."
                
                # Populate based on table structure
                if num_cols == 4:
                    # Simplified hybrid template: Degree | University/Institution | Year | Grade/GPA
                    cells[0].paragraphs[0].clear()
                    degree_text = degree
                    if specialization:
                        degree_text = f"{degree} in {specialization}" if degree else specialization
                    cells[0].paragraphs[0].add_run(degree_text)
                    
                    cells[1].paragraphs[0].clear()
                    institution_text = university or college
                    cells[1].paragraphs[0].add_run(institution_text)
                    
                    cells[2].paragraphs[0].clear()
                    cells[2].paragraphs[0].add_run(str(year))
                    
                    cells[3].paragraphs[0].clear()
                    cells[3].paragraphs[0].add_run(str(grade))
                    
                    print(f"DEBUG: Education {idx+1}: {degree_text} | {institution_text} | {year} | {grade}")
                    
                elif num_cols >= 7:
                    # Full table: Sl. No. | Degree | Branch | Year | College | University | Percentage
                    cells[0].paragraphs[0].clear()
                    cells[0].paragraphs[0].add_run(str(idx + 1))
                    
                    cells[1].paragraphs[0].clear()
                    cells[1].paragraphs[0].add_run(degree)
                    
                    cells[2].paragraphs[0].clear()
                    cells[2].paragraphs[0].add_run(specialization)
                    
                    cells[3].paragraphs[0].clear()
                    cells[3].paragraphs[0].add_run(str(year))
                    
                    cells[4].paragraphs[0].clear()
                    cells[4].paragraphs[0].add_run(college)
                    
                    cells[5].paragraphs[0].clear()
                    cells[5].paragraphs[0].add_run(university)
                    
                    cells[6].paragraphs[0].clear()
                    cells[6].paragraphs[0].add_run(str(grade))
                    
                    print(f"DEBUG: Education {idx+1}: {degree} | {specialization} | {year} | {college} | {university} | {grade}")
                
                row_idx += 1
            
            print(f"DEBUG: Successfully populated education table with {row_idx - 1} entries")
            return True
            
        except Exception as e:
            print(f"Error populating qualification table: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _parse_project_experience(self, project_exp_text: str) -> list:
        """Parse project experience text into structured data"""
        projects = []
        try:
            # Split by project markers
            project_sections = project_exp_text.split("PROJECT:")
            
            for section in project_sections[1:]:  # Skip first empty section
                project = {}
                lines = section.strip().split('\n')
                
                if lines:
                    # First line is project name
                    project["name"] = lines[0].strip()
                    
                    # Parse other details
                    for line in lines[1:]:
                        if "Client:" in line:
                            project["client"] = line.split("Client:", 1)[1].strip()
                        elif "Duration:" in line:
                            project["duration"] = line.split("Duration:", 1)[1].strip()
                        elif "Role:" in line:
                            project["role"] = line.split("Role:", 1)[1].strip()
                        elif "Technologies:" in line:
                            project["technologies"] = line.split("Technologies:", 1)[1].strip()
                        elif "Description:" in line:
                            project["description"] = line.split("Description:", 1)[1].strip()
                
                if project.get("name"):
                    projects.append(project)
                    
        except Exception as e:
            print(f"Error parsing project experience: {e}")
            
        return projects

    def _parse_work_experience(self, work_exp_text) -> list:
        """Parse work experience from structured list or legacy text format."""
        experiences = []
        try:
            if isinstance(work_exp_text, list):
                for exp in work_exp_text:
                    if not isinstance(exp, dict):
                        continue
                    company = exp.get("organization") or exp.get("company") or ""
                    title = exp.get("designation") or exp.get("title") or exp.get("role") or ""
                    start = exp.get("employmentStartDate") or exp.get("startDate") or exp.get("start_date") or ""
                    end = exp.get("employmentEndDate") or exp.get("endDate") or exp.get("end_date") or ""
                    if exp.get("isCurrentCompany") and not end:
                        end = "Present"
                    duration = exp.get("duration") or ""
                    if not duration and (start or end):
                        duration = f"{str(start).strip()} - {str(end).strip()}".strip(" -")

                    parsed = {
                        "title": title,
                        "company": company,
                        "duration": duration,
                        "startDate": start,
                        "endDate": end,
                        "location": exp.get("location") or "",
                        "responsibilities": exp.get("responsibilities") or [],
                    }
                    if parsed.get("title") or parsed.get("company"):
                        experiences.append(parsed)
                return experiences

            # Split by double newlines (experience separators)
            exp_sections = str(work_exp_text).split('\n\n')
            
            for section in exp_sections:
                exp = {}
                lines = section.strip().split('\n')
                
                if lines:
                    # First line usually contains title and company
                    first_line = lines[0]
                    if " at " in first_line:
                        parts = first_line.split(" at ", 1)
                        exp["title"] = parts[0].strip()
                        exp["company"] = parts[1].strip()
                    
                    # Look for duration
                    for line in lines:
                        if "Duration:" in line:
                            exp["duration"] = line.split("Duration:", 1)[1].strip()
                
                if exp.get("title") or exp.get("company"):
                    experiences.append(exp)
                    
        except Exception as e:
            print(f"Error parsing work experience: {e}")
            
        return experiences

    def _parse_certifications(self, cert_text) -> list:
        """Parse certifications from structured list or legacy text format."""
        certifications = []
        try:
            if isinstance(cert_text, list):
                pending_name_idx = -1
                date_like = re.compile(r"^/?\d{1,2}/\d{1,2}/\d{2,4}$|^/?\d{1,2}/\d{2,4}$")
                i = 0
                while i < len(cert_text):
                    cert = cert_text[i]
                    if isinstance(cert, dict):
                        name = cert.get("name") or cert.get("certification") or cert.get("title") or ""
                        if not str(name).strip():
                            i += 1
                            continue
                        certifications.append({
                            "name": str(name).strip(),
                            "year": str(cert.get("year") or cert.get("date") or cert.get("issueDate") or "").strip(),
                            "issuer": str(cert.get("issuer") or cert.get("issuingOrganization") or cert.get("organization") or "").strip(),
                        })
                        pending_name_idx = len(certifications) - 1
                        i += 1
                        continue

                    if isinstance(cert, str) and cert.strip():
                        token = cert.strip().lstrip("•").strip()
                        if not token:
                            i += 1
                            continue

                        if date_like.match(token):
                            if pending_name_idx >= 0 and pending_name_idx < len(certifications):
                                existing = certifications[pending_name_idx].get("year", "").strip()
                                certifications[pending_name_idx]["year"] = f"{existing} - {token}".strip(" -") if existing else token
                            i += 1
                            continue

                        # Merge split name fragments like "SAS" + "ODI tools" when date follows.
                        if (
                            i + 1 < len(cert_text)
                            and isinstance(cert_text[i + 1], str)
                            and cert_text[i + 1].strip()
                            and not date_like.match(cert_text[i + 1].strip().lstrip("•").strip())
                            and i + 2 < len(cert_text)
                            and isinstance(cert_text[i + 2], str)
                            and date_like.match(cert_text[i + 2].strip().lstrip("•").strip())
                        ):
                            token = f"{token}, {cert_text[i + 1].strip().lstrip('•').strip()}"
                            i += 1

                        certifications.append({"name": token, "year": "", "issuer": ""})
                        pending_name_idx = len(certifications) - 1
                    i += 1
                return certifications

            lines = str(cert_text).split('\n')
            for line in lines:
                if line.strip().startswith('•'):
                    cert = {}
                    cert_line = line.strip()[1:].strip()  # Remove bullet
                    
                    # Extract year in parentheses
                    if '(' in cert_line and ')' in cert_line:
                        year_start = cert_line.rfind('(')
                        year_end = cert_line.rfind(')')
                        cert["year"] = cert_line[year_start+1:year_end]
                        cert["name"] = cert_line[:year_start].strip()
                        
                        # Remove issuer part if present
                        if " - " in cert["name"]:
                            cert["name"] = cert["name"].split(" - ")[0].strip()
                    else:
                        cert["name"] = cert_line
                        cert["year"] = ""
                    
                    certifications.append(cert)
                    
        except Exception as e:
            print(f"Error parsing certifications: {e}")
            
        return certifications

    def _parse_education(self, edu_text) -> list:
        """Parse education text into structured data"""
        educations = []
        try:
            # Handle both string and list formats
            if isinstance(edu_text, list):
                # If it's already a list of education objects
                for edu_item in edu_text:
                    if isinstance(edu_item, dict):
                        educations.append(edu_item)
                return educations
            elif isinstance(edu_text, str):
                lines = edu_text.split('\n')
            else:
                return educations
            for line in lines:
                if line.strip().startswith('•'):
                    edu = {}
                    edu_line = line.strip()[1:].strip()  # Remove bullet
                    
                    # Extract year in parentheses
                    if '(' in edu_line and ')' in edu_line:
                        year_start = edu_line.rfind('(')
                        year_end = edu_line.rfind(')')
                        edu["year"] = edu_line[year_start+1:year_end]
                        degree_part = edu_line[:year_start].strip()
                    else:
                        degree_part = edu_line
                        edu["year"] = ""
                    
                    # Parse degree and institution
                    if " from " in degree_part:
                        parts = degree_part.split(" from ", 1)
                        edu["qualification"] = parts[0].strip()
                        edu["college"] = parts[1].strip()
                    elif " in " in degree_part:
                        parts = degree_part.split(" in ", 1)
                        edu["qualification"] = parts[0].strip()
                        edu["specialization"] = parts[1].strip()
                    else:
                        edu["qualification"] = degree_part
                    
                    educations.append(edu)
                    
        except Exception as e:
            print(f"Error parsing education: {e}")
            
        return educations

    def _format_technologies(self, technologies) -> str:
        """Format technologies list for table display"""
        if isinstance(technologies, list):
            return ", ".join(str(tech) for tech in technologies)
        elif isinstance(technologies, str):
            return technologies
        else:
            return ""

    def _format_responsibilities(self, responsibilities) -> str:
        """Format responsibilities list for table display"""
        if isinstance(responsibilities, list):
            # Truncate long responsibility lists for table display
            formatted = []
            for resp in responsibilities[:3]:  # Take first 3 responsibilities
                resp_text = str(resp).strip()
                if resp_text:
                    if len(resp_text) > 60:
                        resp_text = resp_text[:57] + "..."
                    formatted.append(resp_text)
            return "; ".join(formatted)
        elif isinstance(responsibilities, str):
            return responsibilities
        else:
            return ""

    def _format_date_for_table(self, date_str: str) -> str:
        """Format date string for table display (convert 'Mar 2023' to '03/23')"""
        try:
            # Handle common date formats
            date_str = date_str.strip()
            
            # Month name mapping
            month_mapping = {
                'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04',
                'may': '05', 'jun': '06', 'jul': '07', 'aug': '08',
                'sep': '09', 'oct': '10', 'nov': '11', 'dec': '12',
                'january': '01', 'february': '02', 'march': '03', 'april': '04',
                'june': '06', 'july': '07', 'august': '08', 'september': '09',
                'october': '10', 'november': '11', 'december': '12'
            }
            
            # Try to parse "Mar 2023" format
            parts = date_str.lower().split()
            if len(parts) >= 2:
                month_str = parts[0]
                year_str = parts[1]
                
                # Convert month name to number
                for month_name, month_num in month_mapping.items():
                    if month_str.startswith(month_name):
                        # Convert year to 2-digit format
                        if len(year_str) == 4:
                            year_short = year_str[2:]
                        else:
                            year_short = year_str
                        return f"{month_num}/{year_short}"
            
            # If parsing fails, return original
            return date_str
            
        except Exception:
            return date_str

    def _hide_empty_sections(self, doc: Document, context: Dict[str, Any]) -> None:
        """
        Aggressively hide ALL sections with empty content across all template types.
        This works for standard, modern, and hybrid templates.
        """
        try:
            # Comprehensive section mappings: (heading_text_variants, context_keys_to_check)
            section_mappings = [
                # Core sections
                (["PROFESSIONAL SUMMARY", "EXPERIENCE SUMMARY", "SUMMARY"], 
                 ["summary", "professional_summary"]),
                
                # Skills sections
                (["TECHNICAL SKILLS", "PRIMARY SKILLS", "SKILLS"], 
                 ["technical_skills_section", "skills", "primary_skills"]),
                (["CORE COMPETENCIES", "COMPETENCIES"], 
                 ["core_competencies", "skills"]),
                (["SECONDARY SKILLS", "SECONDARY SKILL"], 
                 ["secondary_skills"]),
                (["AI FRAMEWORKS", "AI/ML FRAMEWORKS"], 
                 ["ai_frameworks"]),
                (["CLOUD PLATFORMS", "CLOUD TECHNOLOGIES"], 
                 ["cloud_platforms"]),
                (["TOOLS & PLATFORMS", "DEVELOPMENT TOOLS", "TOOLS AND PLATFORMS"], 
                 ["tools_and_platforms"]),
                (["DATABASES", "DATABASE CONNECTIVITY"], 
                 ["databases"]),
                (["OPERATING SYSTEMS", "OS"], 
                 ["operating_systems"]),
                (["DOMAIN EXPERTISE", "DOMAIN KNOWLEDGE"], 
                 ["domain_expertise"]),
                
                # Experience sections
                (["WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EMPLOYMENT HISTORY"], 
                 ["experience_section", "work_experience"]),
                (["PROJECT EXPERIENCE", "PROJECTS", "KEY PROJECTS"], 
                 ["projects_section", "project_experience"]),
                
                # Education & Credentials
                (["EDUCATION", "QUALIFICATION DETAILS", "ACADEMIC BACKGROUND"], 
                 ["education_section", "education"]),
                (["CERTIFICATIONS", "TRAINING", "CERTIFICATIONS & TRAINING"], 
                 ["certifications_section", "certifications"]),
                
                # Optional sections
                (["LEADERSHIP & MENTORING", "LEADERSHIP", "LEADERSHIP & IMPACT"], 
                 ["leadership_lines", "leadership"]),
                (["KEY ACHIEVEMENTS", "ACHIEVEMENTS"], 
                 ["key_achievements_section", "key_achievements"]),
                (["AWARDS & RECOGNITION", "AWARDS", "RECOGNITION"], 
                 ["awards"]),
                (["PUBLICATIONS & PRESENTATIONS", "PUBLICATIONS", "PRESENTATIONS"], 
                 ["publications"]),
                (["LANGUAGES"], 
                 ["languages"]),
            ]
            
            # Track elements to delete (paragraphs and tables)
            elements_to_delete = []
            
            # Iterate through all document elements
            for i, paragraph in enumerate(doc.paragraphs):
                paragraph_text = paragraph.text.strip().upper()
                
                # Check if this paragraph matches any section heading
                for heading_variants, context_keys in section_mappings:
                    # Check if paragraph matches any heading variant
                    is_section_heading = any(
                        paragraph_text == variant.upper() 
                        for variant in heading_variants
                    )
                    
                    if is_section_heading:
                        # Check if ANY of the mapped context keys has content
                        has_content = False
                        for context_key in context_keys:
                            content = context.get(context_key)
                            
                            # Comprehensive empty check
                            if content is not None:
                                if isinstance(content, str) and content.strip():
                                    has_content = True
                                    break
                                elif isinstance(content, list) and len(content) > 0:
                                    # Check if list has any non-empty items
                                    if any(item for item in content if item):
                                        has_content = True
                                        break
                                elif isinstance(content, dict) and len(content) > 0:
                                    has_content = True
                                    break
                                elif isinstance(content, (int, float, bool)):
                                    has_content = True
                                    break
                        
                        # If section is empty, mark for deletion
                        if not has_content:
                            # Mark the heading paragraph
                            elements_to_delete.append(paragraph)
                            
                            # Check what follows the heading and mark for deletion
                            # It could be: content paragraph(s), table, or both
                            next_idx = i + 1
                            
                            # Look ahead to find related content
                            while next_idx < len(doc.paragraphs):
                                next_para = doc.paragraphs[next_idx]
                                next_text = next_para.text.strip().upper()
                                
                                # Stop if we hit another section heading
                                is_next_section = False
                                for other_variants, _ in section_mappings:
                                    if any(next_text == v.upper() for v in other_variants):
                                        is_next_section = True
                                        break
                                
                                if is_next_section:
                                    break
                                
                                # Mark content paragraphs for deletion
                                # Stop at first non-empty paragraph that's not a placeholder
                                if next_text and not next_text.startswith("{{") and not next_text.endswith("}}"):
                                    # This might be content, but check if it's actually empty
                                    if not next_text or next_text in ["", " ", "\n"]:
                                        elements_to_delete.append(next_para)
                                        next_idx += 1
                                    else:
                                        # Found real content, don't delete more
                                        break
                                else:
                                    # Empty or placeholder paragraph
                                    elements_to_delete.append(next_para)
                                    next_idx += 1
                                
                                # Safety: don't look too far ahead
                                if next_idx - i > 10:
                                    break
                            
                            # Also check for tables that follow this heading
                            # Tables are separate from paragraphs in python-docx
                            para_element = paragraph._element
                            next_element = para_element.getnext()
                            
                            # Check a few next elements for tables
                            check_count = 0
                            while next_element is not None and check_count < 5:
                                # Check if it's a table
                                if next_element.tag.endswith('}tbl'):
                                    # Mark table for deletion
                                    elements_to_delete.append(next_element)
                                    break
                                # Stop if we hit another paragraph with section heading
                                elif next_element.tag.endswith('}p'):
                                    # Check if this is a new section
                                    temp_para_text = ""
                                    for text_elem in next_element.iter():
                                        if text_elem.tag.endswith('}t'):
                                            temp_para_text += text_elem.text or ""
                                    
                                    temp_para_text = temp_para_text.strip().upper()
                                    is_new_section = any(
                                        any(temp_para_text == v.upper() for v in variants)
                                        for variants, _ in section_mappings
                                    )
                                    
                                    if is_new_section:
                                        break
                                
                                next_element = next_element.getnext()
                                check_count += 1
                        
                        break  # Found matching section, no need to check other variants
            
            # Delete marked elements (from bottom to top to avoid index issues)
            deleted_count = 0
            for element in reversed(elements_to_delete):
                try:
                    if hasattr(element, '_element'):
                        # It's a paragraph wrapper
                        p = element._element
                        parent = p.getparent()
                        if parent is not None:
                            parent.remove(p)
                            deleted_count += 1
                    else:
                        # It's a raw element (like table)
                        parent = element.getparent()
                        if parent is not None:
                            parent.remove(element)
                            deleted_count += 1
                except Exception as e:
                    print(f"Warning: Could not delete element: {e}")
            
            if deleted_count > 0:
                print(f"DEBUG: Aggressively hid {deleted_count} empty section elements")
            else:
                print("DEBUG: No empty sections found to hide")
            
        except Exception as e:
            print(f"Warning: Could not hide empty sections: {e}")
            import traceback
            traceback.print_exc()
    
    def format_projects_modern(self, projects_data: list) -> str:
        """
        Format projects for modern template (2026 best practices)
        Returns clean, formatted text without complex tables
        """
        if not projects_data:
            return ""
        
        formatted_sections = []
        
        for project in projects_data:
            if isinstance(project, dict):
                lines = []
                
                # Project Name (bold heading)
                project_name = project.get("project_name", project.get("name", "Unnamed Project"))
                lines.append(f"**{project_name}**")
                
                # Metadata line (Client, Duration, Role)
                metadata_parts = []
                if project.get("client"):
                    metadata_parts.append(f"Client: {project['client']}")
                if project.get("duration"):
                    metadata_parts.append(f"Duration: {project['duration']}")
                if project.get("role", project.get("position")):
                    role = project.get("role", project.get("position"))
                    metadata_parts.append(f"Role: {role}")
                
                if metadata_parts:
                    lines.append(" | ".join(metadata_parts))
                
                # Description
                description = project.get("project_description", project.get("description", ""))
                if description:
                    lines.append(f"\n{description}")
                
                # Technologies
                technologies = self._format_technologies(
                    project.get("technologies_used", 
                    project.get("technologies", 
                    project.get("tech_stack", [])))
                )
                if technologies:
                    lines.append(f"\nTechnologies: {technologies}")
                
                # Key Contributions
                contributions = project.get("responsibilities", 
                                           project.get("key_responsibilities", 
                                           project.get("contributions", [])))
                if contributions:
                    lines.append("\nKey Contributions:")
                    if isinstance(contributions, list):
                        for contrib in contributions:
                            lines.append(f"• {contrib}")
                    elif isinstance(contributions, str):
                        lines.append(f"• {contributions}")
                
                formatted_sections.append("\n".join(lines))
        
        return "\n\n".join(formatted_sections)
    
    def format_experience_modern(self, experience_data: list) -> str:
        """
        Format work experience for modern template (2026 best practices)
        Returns clean timeline format
        """
        if not experience_data:
            return ""
        
        formatted_sections = []
        
        for exp in experience_data:
            if isinstance(exp, dict):
                lines = []
                
                # Company & Location
                company = exp.get("company", exp.get("organization", ""))
                location = exp.get("location", "")
                if company:
                    company_line = company
                    if location:
                        company_line += f" • {location}"
                    lines.append(f"**{company_line}**")
                
                # Role & Duration
                role = exp.get("title", exp.get("designation", exp.get("role", "")))
                duration = exp.get("duration", "")
                if role:
                    role_line = role
                    if duration:
                        role_line += f" | {duration}"
                    lines.append(role_line)
                
                # Achievements/Responsibilities
                responsibilities = exp.get("responsibilities", 
                                          exp.get("achievements", []))
                if responsibilities:
                    if isinstance(responsibilities, list):
                        for resp in responsibilities:
                            lines.append(f"• {resp}")
                    elif isinstance(responsibilities, str):
                        lines.append(f"• {responsibilities}")
                
                formatted_sections.append("\n".join(lines))
        
        return "\n\n".join(formatted_sections)
    
    def format_skills_matrix(self, context: Dict[str, Any]) -> str:
        """
        Format skills in clean category: technologies format
        Following 2026 best practices for skill presentation
        """
        skills_categories = [
            ("Primary Skills", context.get("skills", "")),
            ("Secondary Skills", context.get("secondary_skills", "")),
            ("AI/ML Frameworks", context.get("ai_frameworks", "")),
            ("Cloud Platforms", context.get("cloud_platforms", "")),
            ("Development Tools", context.get("tools_and_platforms", "")),
            ("Databases", context.get("databases", "")),
            ("Operating Systems", context.get("operating_systems", "")),
            ("Domain Expertise", context.get("domain_expertise", "")),
        ]
        
        formatted_lines = []
        for category, skills in skills_categories:
            if skills:
                # Format skills if they're a list
                if isinstance(skills, list):
                    skills_str = ", ".join(str(s) for s in skills if s)
                else:
                    skills_str = str(skills)
                
                if skills_str.strip():
                    formatted_lines.append(f"{category}: {skills_str}")
        
        return "\n".join(formatted_lines)

    def _apply_modern_enhancements(self, doc: Document, context: Dict[str, Any]) -> None:
        """Apply modern template enhancements - minimal tables, clean typography"""
        try:
            # Modern-specific enhancements (empty sections already handled globally)
            print("DEBUG: Applied modern template enhancements")
        except Exception as e:
            print(f"Warning: Could not apply modern enhancements: {e}")
    
    def _apply_hybrid_enhancements(self, doc: Document, context: Dict[str, Any]) -> None:
        """Apply hybrid template enhancements - keep skills tables, clean experience format"""
        try:
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Hybrid approach: selective enhancements
            # 1. Keep skills tables structured (good for scanning)
            # 2. Use cleaner formatting for experience/projects
            # 3. Hide truly empty optional sections only
            
            hidden_count = 0
            
            # Hide only truly optional empty sections
            optional_sections = [
                ("AWARDS & RECOGNITION", "awards"),
                ("PUBLICATIONS & PRESENTATIONS", "publications"),
            ]
            
            paragraphs_to_delete = []
            for i, paragraph in enumerate(doc.paragraphs):
                paragraph_text = paragraph.text.strip()
                
                for heading_text, context_key in optional_sections:
                    if paragraph_text == heading_text:
                        content = context.get(context_key, "")
                        if not content or content.strip() == "":
                            paragraphs_to_delete.append(paragraph)
                            if i + 1 < len(doc.paragraphs):
                                paragraphs_to_delete.append(doc.paragraphs[i + 1])
                            hidden_count += 1
                        break
            
            # Delete marked paragraphs
            for paragraph in paragraphs_to_delete:
                p = paragraph._element
                p.getparent().remove(p)
            
            if hidden_count > 0:
                print(f"DEBUG: Hid {hidden_count} empty optional sections in hybrid template")
            
            print("DEBUG: Applied hybrid template enhancements")
        except Exception as e:
            print(f"Warning: Could not apply hybrid enhancements: {e}")

    def _replace_table_with_clean_projects(self, table, context: Dict[str, Any]) -> bool:
        """
        Replace project table with clean formatted text (2026 best practices)
        For hybrid/modern templates
        """
        try:
            # Get project experience data
            project_exp = context.get("project_experience", "")
            projects = []
            
            if isinstance(project_exp, str):
                projects = self._parse_project_experience(project_exp)
            elif isinstance(project_exp, list):
                for proj in project_exp:
                    if isinstance(proj, dict):
                        # Parse description to extract client and actual description
                        description = proj.get("project_description", proj.get("description", ""))
                        client = proj.get("client", "")
                        
                        # Handle markdown-formatted descriptions like "**Client:** X **Description:** Y"
                        if "**Client:**" in description or "**Description:**" in description:
                            # Extract client
                            if "**Client:**" in description:
                                client_match = re.search(r'\*\*Client:\*\*\s*([^\*]+)', description)
                                if client_match:
                                    client = client_match.group(1).strip()
                            
                            # Extract actual description
                            if "**Description:**" in description:
                                desc_match = re.search(r'\*\*Description:\*\*\s*(.+)', description)
                                if desc_match:
                                    description = desc_match.group(1).strip()
                        
                        mapped_project = {
                            "name": proj.get("project_name", proj.get("projectName", proj.get("name", ""))),
                            "client": proj.get("client_name", proj.get("clientName", client)),
                            "description": proj.get("projectDescription", description),
                            "technologies": self._format_technologies(
                                proj.get("technologies_used", 
                                proj.get("technologies", 
                                proj.get("tech_stack", proj.get("toolsUsed", proj.get("environment", [])))))
                            ),
                            "duration": proj.get("duration", "") or self._build_project_duration(proj),
                            "role": proj.get("role", proj.get("position", "")),
                            "contributions": proj.get("responsibilities", 
                                proj.get("key_responsibilities", 
                                proj.get("contributions", []))),
                            "team_size": proj.get("team_size", proj.get("teamSize", ""))
                        }
                        projects.append(mapped_project)
            
            if not projects:
                return False
            
            # Get the parent element of the table
            table_element = table._element
            parent = table_element.getparent()
            table_index = parent.index(table_element)
            
            # Remove the table
            parent.remove(table_element)
            
            # Insert formatted project sections using proper Word formatting
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            insert_pos = table_index
            
            for i, project in enumerate(projects):
                # Project Name - Bold heading
                p_name = OxmlElement('w:p')
                pPr_name = OxmlElement('w:pPr')
                p_name.append(pPr_name)
                
                r_name = OxmlElement('w:r')
                rPr_name = OxmlElement('w:rPr')
                b_name = OxmlElement('w:b')
                sz_name = OxmlElement('w:sz')
                sz_name.set(qn('w:val'), '24')  # 12pt font
                rPr_name.append(b_name)
                rPr_name.append(sz_name)
                r_name.append(rPr_name)
                
                t_name = OxmlElement('w:t')
                project_name = project.get("name", "Unnamed Project")
                # Remove "PROJECT:" prefix if present
                if project_name.upper().startswith("PROJECT:"):
                    project_name = project_name[8:].strip()
                t_name.text = project_name.upper()
                r_name.append(t_name)
                p_name.append(r_name)
                
                parent.insert(insert_pos, p_name)
                insert_pos += 1
                
                # Empty line for spacing
                p_space1 = OxmlElement('w:p')
                parent.insert(insert_pos, p_space1)
                insert_pos += 1
                
                # Client, Duration, Role on separate labeled lines
                if project.get("client"):
                    p_client = OxmlElement('w:p')
                    # Bold label
                    r_client_label = OxmlElement('w:r')
                    rPr_client_label = OxmlElement('w:rPr')
                    b_client_label = OxmlElement('w:b')
                    rPr_client_label.append(b_client_label)
                    r_client_label.append(rPr_client_label)
                    t_client_label = OxmlElement('w:t')
                    t_client_label.text = "Client: "
                    r_client_label.append(t_client_label)
                    p_client.append(r_client_label)
                    # Regular text
                    r_client_val = OxmlElement('w:r')
                    t_client_val = OxmlElement('w:t')
                    t_client_val.text = project['client']
                    r_client_val.append(t_client_val)
                    p_client.append(r_client_val)
                    parent.insert(insert_pos, p_client)
                    insert_pos += 1
                
                # Description with bold label
                if project.get("description"):
                    p_desc = OxmlElement('w:p')
                    # Bold label
                    r_desc_label = OxmlElement('w:r')
                    rPr_desc_label = OxmlElement('w:rPr')
                    b_desc_label = OxmlElement('w:b')
                    rPr_desc_label.append(b_desc_label)
                    r_desc_label.append(rPr_desc_label)
                    t_desc_label = OxmlElement('w:t')
                    t_desc_label.text = "Description: "
                    r_desc_label.append(t_desc_label)
                    p_desc.append(r_desc_label)
                    # Regular text
                    r_desc_val = OxmlElement('w:r')
                    t_desc_val = OxmlElement('w:t')
                    t_desc_val.text = project["description"]
                    r_desc_val.append(t_desc_val)
                    p_desc.append(r_desc_val)
                    parent.insert(insert_pos, p_desc)
                    insert_pos += 1
                
                # Roles and Responsibilities with bold label
                contributions = project.get("contributions", [])
                if contributions:
                    p_roles = OxmlElement('w:p')
                    r_roles_label = OxmlElement('w:r')
                    rPr_roles_label = OxmlElement('w:rPr')
                    b_roles_label = OxmlElement('w:b')
                    rPr_roles_label.append(b_roles_label)
                    r_roles_label.append(rPr_roles_label)
                    t_roles_label = OxmlElement('w:t')
                    t_roles_label.text = "Roles and Responsibilities:"
                    r_roles_label.append(t_roles_label)
                    p_roles.append(r_roles_label)
                    parent.insert(insert_pos, p_roles)
                    insert_pos += 1
                    
                    if isinstance(contributions, list):
                        for contrib in contributions:
                            p_contrib = OxmlElement('w:p')
                            # Add bullet using symbol
                            r_contrib = OxmlElement('w:r')
                            t_contrib = OxmlElement('w:t')
                            t_contrib.set(qn('xml:space'), 'preserve')
                            t_contrib.text = f"  • {contrib}"
                            r_contrib.append(t_contrib)
                            p_contrib.append(r_contrib)
                            parent.insert(insert_pos, p_contrib)
                            insert_pos += 1
                    elif isinstance(contributions, str):
                        p_contrib = OxmlElement('w:p')
                        r_contrib = OxmlElement('w:r')
                        t_contrib = OxmlElement('w:t')
                        t_contrib.set(qn('xml:space'), 'preserve')
                        t_contrib.text = f"  • {contributions}"
                        r_contrib.append(t_contrib)
                        p_contrib.append(r_contrib)
                        parent.insert(insert_pos, p_contrib)
                        insert_pos += 1
                
                # Technologies with bold label
                if project.get("technologies"):
                    p_tech = OxmlElement('w:p')
                    # Bold label
                    r_tech_label = OxmlElement('w:r')
                    rPr_tech_label = OxmlElement('w:rPr')
                    b_tech_label = OxmlElement('w:b')
                    rPr_tech_label.append(b_tech_label)
                    r_tech_label.append(rPr_tech_label)
                    t_tech_label = OxmlElement('w:t')
                    t_tech_label.text = "Technologies: "
                    r_tech_label.append(t_tech_label)
                    p_tech.append(r_tech_label)
                    # Regular text
                    r_tech_val = OxmlElement('w:r')
                    t_tech_val = OxmlElement('w:t')
                    t_tech_val.text = project["technologies"]
                    r_tech_val.append(t_tech_val)
                    p_tech.append(r_tech_val)
                    parent.insert(insert_pos, p_tech)
                    insert_pos += 1
                
                # Add spacing between projects
                p_space2 = OxmlElement('w:p')
                parent.insert(insert_pos, p_space2)
                insert_pos += 1
            
            print(f"DEBUG: Replaced project table with clean format ({len(projects)} projects)")
            return True
            
        except Exception as e:
            print(f"Error replacing project table: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _replace_table_with_clean_experience(self, table, context: Dict[str, Any]) -> bool:
        """
        Replace experience table with timeline format (2026 best practices)
        For hybrid/modern templates
        """
        try:
            # Get work experience data
            work_exp = context.get("work_experience", "")
            if not work_exp:
                return False
                
            experiences = self._parse_work_experience(work_exp)
            if not experiences:
                return False
            
            # Get the parent element of the table
            table_element = table._element
            parent = table_element.getparent()
            table_index = parent.index(table_element)
            
            # Remove the table
            parent.remove(table_element)
            
            # Insert formatted experience sections
            from docx.oxml import OxmlElement
            
            for i, exp in enumerate(experiences):
                # Company & Location - Bold
                p1 = OxmlElement('w:p')
                pPr1 = OxmlElement('w:pPr')
                p1.append(pPr1)
                
                r1 = OxmlElement('w:r')
                rPr1 = OxmlElement('w:rPr')
                b1 = OxmlElement('w:b')
                rPr1.append(b1)
                r1.append(rPr1)
                
                t1 = OxmlElement('w:t')
                company = exp.get("company", "")
                location = exp.get("location", "")
                if location:
                    t1.text = f"{company} • {location}"
                else:
                    t1.text = company
                r1.append(t1)
                p1.append(r1)
                
                parent.insert(table_index + (i * 4), p1)
                
                # Role & Duration
                p2 = OxmlElement('w:p')
                r2 = OxmlElement('w:r')
                t2 = OxmlElement('w:t')
                role = exp.get("title", "")
                duration = exp.get("duration", "")
                if duration:
                    t2.text = f"{role} | {duration}"
                else:
                    t2.text = role
                r2.append(t2)
                p2.append(r2)
                
                parent.insert(table_index + (i * 4) + 1, p2)
                
                # Responsibilities (if available)
                responsibilities = exp.get("responsibilities", [])
                if responsibilities:
                    if isinstance(responsibilities, list):
                        for j, resp in enumerate(responsibilities):
                            p_resp = OxmlElement('w:p')
                            r_resp = OxmlElement('w:r')
                            t_resp = OxmlElement('w:t')
                            t_resp.text = f"• {resp}"
                            r_resp.append(t_resp)
                            p_resp.append(r_resp)
                            parent.insert(table_index + (i * 4) + 2 + j, p_resp)
                
                # Add spacing paragraph
                p_space = OxmlElement('w:p')
                parent.insert(table_index + (i * 4) + 3, p_space)
            
            print(f"DEBUG: Replaced experience table with timeline format ({len(experiences)} experiences)")
            return True
            
        except Exception as e:
            print(f"Error replacing experience table: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _replace_table_with_bullet_certifications(self, table, context: Dict[str, Any]) -> bool:
        """
        Replace certification table with clean bullet list (2026 best practices)
        For hybrid/modern templates
        """
        try:
            certifications = context.get("certifications", "")
            if not certifications:
                return False
                
            certs = self._parse_certifications(certifications)
            if not certs:
                return False
            
            # Get the parent element of the table
            table_element = table._element
            parent = table_element.getparent()
            table_index = parent.index(table_element)
            
            # Remove the table
            parent.remove(table_element)
            
            # Insert formatted certifications as bullets
            from docx.oxml import OxmlElement
            
            for i, cert in enumerate(certs):
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                
                cert_text = cert.get("name", "")
                year = cert.get("year", "")
                issuer = cert.get("issuer", "")
                parts = [p for p in [issuer, year] if str(p).strip()]
                suffix = f" ({' | '.join(parts)})" if parts else ""
                t.text = f"• {cert_text}{suffix}"
                
                r.append(t)
                p.append(r)
                parent.insert(table_index + i, p)
            
            print(f"DEBUG: Replaced certification table with bullet list ({len(certs)} certifications)")
            return True
            
        except Exception as e:
            print(f"Error replacing certification table: {e}")
            return False

    def _split_summary_lines(self, summary_value: Any) -> list:
        raw = str(summary_value or "").replace("\r\n", "\n").strip()
        if not raw:
            return []

        normalized = re.sub(r"[\u2022\u25E6\u25AA\u25CF\u25C6\u25BA\uF076]", "\n• ", raw)
        lines = [line.strip() for line in normalized.split("\n") if line.strip()]

        if len(lines) == 1:
            single = lines[0]
            single = re.sub(r"\s+([a-zA-Z]\.)\s+", r"\n\1 ", single)
            single = re.sub(r"\s+(\d+\.)\s+", r"\n\1 ", single)
            lines = [line.strip() for line in single.split("\n") if line.strip()]

        cleaned = []
        seen = set()
        for line in lines:
            cleaned_line = re.sub(r"^([\-*\u2022•]|\d+[.)]|[a-zA-Z][.)])\s+", "", line).strip()
            cleaned_line = self._dedupe_experience_phrase(cleaned_line)
            if cleaned_line:
                key = cleaned_line.lower()
                if key not in seen:
                    seen.add(key)
                    cleaned.append(cleaned_line)

        return cleaned

    def _dedupe_experience_phrase(self, text: str) -> str:
        """Remove duplicate experience phrases like repeated '10+ years' in one line."""
        if not text:
            return ""
        pattern = re.compile(r"(?i)(\b\d+(?:\.\d+)?\s*\+?\s*years?\b|\b\d+\s*\+\s*years?\b)")
        matches = list(pattern.finditer(text))
        if len(matches) <= 1:
            return text.strip()

        first = matches[0]
        keep_parts = [text[:first.end()]]
        cursor = first.end()
        for m in matches[1:]:
            keep_parts.append(text[cursor:m.start()])
            cursor = m.end()
        keep_parts.append(text[cursor:])
        compact = "".join(keep_parts)
        compact = re.sub(r"\s{2,}", " ", compact)
        compact = re.sub(r"\s+,", ",", compact)
        return compact.strip()

    def _build_project_duration(self, project: Dict[str, Any]) -> str:
        start = project.get("durationFrom") or project.get("startDate") or ""
        end = project.get("durationTo") or project.get("endDate") or ""
        if project.get("isCurrentProject") and not end:
            end = "Present"
        if start or end:
            return f"{str(start).strip()} - {str(end).strip()}".strip(" -")
        return ""

    def _render_fallback(self, context: Dict[str, Any]) -> bytes:
        # Create a simple text-based document representation
        lines = []
        lines.append("=== NTT DATA PROFESSIONAL CV ===\n")
        
        # Personal Information
        lines.append("PERSONAL INFORMATION:")
        info_fields = [
            ("Name", "full_name"),
            ("Employee ID", "employee_id"),
            ("Email", "email"),
            ("Contact", "contact_number"),
            ("Title", "current_title"),
            ("Grade", "grade"),
            ("Location", "location"),
            ("Organization", "organization"),
            ("Experience", "experience"),
            ("Target Role", "target_role"),
        ]
        
        for label, key in info_fields:
            value = context.get(key, "")
            if value:
                lines.append(f"{label}: {value}")
        
        lines.append("")
        
        # Content sections
        content_sections = [
            ("PROFESSIONAL SUMMARY", "summary"),
            ("SKILLS", "skills"),
            ("SECONDARY SKILLS", "secondary_skills"),
            ("TOOLS & PLATFORMS", "tools_and_platforms"),
            ("WORK EXPERIENCE", "work_experience"),
            ("PROJECT EXPERIENCE", "project_experience"),
            ("EDUCATION", "education"),
            ("CERTIFICATIONS", "certifications"),
            ("LANGUAGES", "languages"),
            ("AWARDS", "awards"),
            ("PUBLICATIONS", "publications"),
        ]
        
        for section_title, key in content_sections:
            content = context.get(key, "")
            if content:
                lines.append(f"{section_title}:")
                if key == "summary":
                    summary_lines = self._split_summary_lines(content)
                    if len(summary_lines) > 1:
                        for line in summary_lines:
                            lines.append(f"- {line}")
                    else:
                        lines.append(str(content))
                else:
                    lines.append(str(content))
                lines.append("")
        
        # Convert to bytes (simplified fallback)
        content_str = "\n".join(lines)
        return content_str.encode('utf-8')
