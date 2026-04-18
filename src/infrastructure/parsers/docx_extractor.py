import html
import re
from zipfile import ZipFile


_RICH_SECTION_MARKERS = (
    "education",
    "qualification details",
    "project details",
    "experience details",
    "training attended",
    "certifications",
)


def _extract_docx_ordered_text(file_path: str) -> str:
    """Extract text by traversing docx paragraph/table blocks in body order."""
    from docx import Document

    doc = Document(file_path)
    full_text = []

    for element in doc.element.body:
        if element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == element:
                    if para.text.strip():
                        full_text.append(para.text)
                    break
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            full_text.append(" | ".join(row_text))
                    break

    return "\n".join(full_text)


def _extract_docx_xml_fallback_text(file_path: str) -> str:
    """Extract text directly from word/document.xml to capture content in structured document blocks."""
    with ZipFile(file_path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")

    # Preserve paragraph/row boundaries before stripping tags.
    xml = re.sub(r"</w:tr>", "\n", xml)
    xml = re.sub(r"</w:p>", "\n", xml)
    xml = re.sub(r"</w:tc>", " | ", xml)
    text = re.sub(r"<[^>]+>", "", xml)
    text = html.unescape(text)

    lines = []
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if line:
            lines.append(line)

    return "\n".join(lines)


def _is_richer_text(candidate: str, baseline: str) -> bool:
    candidate_lower = (candidate or "").lower()
    baseline_lower = (baseline or "").lower()

    baseline_hits = sum(1 for marker in _RICH_SECTION_MARKERS if marker in baseline_lower)
    candidate_hits = sum(1 for marker in _RICH_SECTION_MARKERS if marker in candidate_lower)

    if candidate_hits > baseline_hits:
        return True
    if len(candidate or "") > int(len(baseline or "") * 1.30):
        return True
    return False

def extract_docx(file_path: str) -> str:
    """Extract text from a DOCX file including tables. Requires python-docx"""
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("Install python-docx") from e

    # Validate DOCX parse upfront and then use robust extraction pipeline.
    Document(file_path)

    primary_text = _extract_docx_ordered_text(file_path)
    fallback_text = _extract_docx_xml_fallback_text(file_path)

    if _is_richer_text(fallback_text, primary_text):
        return fallback_text

    return primary_text
